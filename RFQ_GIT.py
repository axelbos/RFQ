import sys
import os
import shutil
import xml.etree.ElementTree as ET
from docx import Document
from docx.oxml.ns import qn
from docxcompose.composer import Composer
import pandas as pd
import re
from collections import defaultdict
from datetime import datetime
from docx.enum.section import WD_SECTION
from docx.shared import Pt

# --- Funktioner ---

def normalize_key(text):
    text = text.lower().strip()
    text = re.sub(r"[ /,\[\];:()\-]", "_", text)
    text = re.sub(r"_+", "_", text)
    return text.strip("_")

def translate_value_if_possible(value, key=None):
    if key == "numpag":
        return "{{numpag}}"  # Låt stå kvar tills vi byter ut den i generate_final_doc
    norm_key = normalize_key(value)
    if key == "counterweight_with_safety_gear":
        return "Ja" if value.strip() == "1" else "Nej"
    result = translation_dict.get(norm_key, value)
    print(f" Placeholder: '{key}' - '{value}' - '{result}'")
    return result

def is_valid_elevator(data):
    return "general_information" in data and data["general_information"].strip()

def extract_machineroom_type(desc):
    match = re.search(r"[PGB]([WTSU])", desc.upper())
    return match.group(1) if match else ""

def extract_multiple_elevators(xml_path):
    tree = ET.parse(xml_path)
    root = tree.getroot()
    ns = {'ns': root.tag.split('}')[0].strip('{')} if '}' in root.tag else {}

    global_data = {}
    grouped_elevators = []
    manual_elevators = []
    current_elevator = {}
    last_general_info = None
    manual_hiss_skapad = False
    grouped_buffer = []
    current_group_mode = False
    expected_group_columns = None

    for table in root.findall('.//Table', ns):
        rows = table.findall('TR', ns)
        if not rows:
            continue

        header_cells = [c for c in rows[0] if c.tag in ('TH', 'TD')]
        num_columns = len(header_cells)
        header_texts = [c.text.strip() if c.text else '' for c in header_cells]
        print(f" Tabell med {num_columns} kolumner: {header_texts}")

        treat_as_grouped = num_columns > 2

        if treat_as_grouped:
            if not current_group_mode:
                expected_group_columns = num_columns
                num_hissar = num_columns - 1
                grouped_buffer = [{} for _ in range(num_hissar)]
                current_group_mode = True
            elif num_columns != expected_group_columns:
                for e in grouped_buffer:
                    if is_valid_elevator(e):
                        grouped_elevators.append(e)
                grouped_buffer = []
                current_group_mode = False
                expected_group_columns = None
                continue

            for row in rows:
                cells = [c for c in row if c.tag in ('TH', 'TD')]
                if len(cells) >= num_columns:
                    key = normalize_key((cells[0].text or '').strip())
                    if not key:
                        continue
                    for i in range(1, num_columns):
                        value = cells[i].text.strip() if cells[i].text else ""
                        grouped_buffer[i - 1][key] = value
                        grouped_buffer[i - 1]["group_id"] = f"group_{len(grouped_elevators) // (num_columns - 1) + 1}"

        elif num_columns in (1, 2):
            temp_data = {}
            for row in rows:
                cells = [c for c in row if c.tag in ('TH', 'TD')]
                if len(cells) == 1:
                    key = normalize_key((cells[0].text or '').strip())
                    if key:
                        temp_data[key] = ""
                elif len(cells) == 2:
                    key = normalize_key((cells[0].text or '').strip())
                    value = (cells[1].text or '').strip()
                    if key:
                        temp_data[key] = value

            general_info = temp_data.get("general_information")
            if general_info:
                if general_info != last_general_info:
                    if is_valid_elevator(current_elevator):
                        manual_elevators.append(current_elevator)
                    print(f" Startar ny singelhiss: {general_info}")
                    current_elevator = temp_data
                    last_general_info = general_info
                    manual_hiss_skapad = True
                else:
                    current_elevator.update(temp_data)
            elif last_general_info:
                current_elevator.update(temp_data)
            else:
                global_data.update(temp_data)

    if current_group_mode:
        for e in grouped_buffer:
            if is_valid_elevator(e):
                grouped_elevators.append(e)

    if is_valid_elevator(current_elevator) and manual_hiss_skapad:
        manual_elevators.append(current_elevator)

    elevators = grouped_elevators + manual_elevators

    print(f"\n{len(elevators)} hissar hittade i XML")
    for idx, e in enumerate(elevators, 1):
        print(f"\n--- HISS {idx} ---")
        for k, v in e.items():
            print(f"{k}: {v}")

    print("\n--- GLOBAL DATA ---")
    for k, v in global_data.items():
        print(f"{k}: {v}")

    return elevators, global_data

def extract_elevator_groups_from_xml(xml_path):
    tree = ET.parse(xml_path)
    root = tree.getroot()
    ns = {'ns': root.tag.split('}')[0].strip('{')} if '}' in root.tag else {}

    groups = []
    for table in root.findall('.//Table', ns):
        rows = table.findall('TR', ns)
        if not rows:
            continue
        header_cells = [c.text.strip() if c.text else "" for c in rows[0] if c.tag in ('TH', 'TD')]
        if not header_cells:
            continue
        if normalize_key(header_cells[0]) == "general_information" and len(header_cells) > 1:
            names = [normalize_key(cell) for cell in header_cells[1:]]
            group = {
                "hissbeteckning": ", ".join(name.split("_")[0].upper() for name in names if name),
                "antal_hissar": str(len(names))
            }
            groups.append(group)
    return groups

def group_elevators_by_spec(elevators):
    key_fields = [
        "range_of_use", "rated_load_q_kg", "rated_speed_v_m_s", "number_of_floors",
        "car_entrance_type", "door_type", "car_shell_width_bb_mm", "car_shell_depth_dd_mm",
        "car_clear_intern_height_ch_mm", "shaft_width_ww_mm", "shaft_depth_wd_mm",
        "min_shaft_pit_depth_ph_mm", "shaft_headroom_height_sh_mm",
        "fire_class_country", "counterweight_with_safety_gear", "control_system"
    ]

    groups = defaultdict(list)
    for e in elevators:
        spec_key = tuple((k, e.get(k, "")) for k in key_fields)
        groups[spec_key].append(e)

    grouped = []
    for group in groups.values():
        base = group[0].copy()
        base["antal_hissar"] = str(len(group))
        base["hissbeteckning"] = ", ".join(e.get("general_information", "").split()[0].upper() for e in group if e.get("general_information"))
        desc = base.get("network_description", "")
        mr_type = extract_machineroom_type(desc)
        if mr_type:
            base["machineroom_type"] = mr_type + "2" if len(group) > 1 else mr_type
        grouped.append(base)

    return grouped

def adjust_translation_by_context(paragraph_text: str, placeholder: str, translated: str) -> str:
    """Justera versalisering beroende på meningens position, men skydda koder och förkortningar."""
    try:
        before = paragraph_text.split(placeholder)[0].strip()
        starts_sentence = not before or before.endswith((". ", ": ", "? ", "! ", "\n"))

        # Skydda om översättningen börjar med t.ex. LED-, AISI 441, HR64, KES800 etc
        if re.match(r"^[A-ZÅÄÖ]{2,}([- ]|$)", translated):  # t.ex. LED- eller AISI 
            return translated

        # Skydda hisskoder som A1, A4 osv
        if translated.isupper() or (
            len(translated) >= 2 and translated[:2].isupper() and any(char.isdigit() for char in translated)
        ):
            return translated

        if starts_sentence:
            return translated[0].upper() + translated[1:]
        else:
            return translated[0].lower() + translated[1:]
    except Exception:
        return translated

def fill_placeholders_in_doc(doc, data, suppress_keys: list = None):
    import re
    normalized_data = {normalize_key(k): v for k, v in data.items()}
    pattern = r"\{\{(.*?)\}\}"
    filled_keys = set(normalized_data.keys())
    print("✅ Ifyllda nycklar:", filled_keys)

    if suppress_keys is None:
        suppress_keys = []
    suppress_keys = set(normalize_key(k) for k in suppress_keys)

    def process_runs(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        matches = re.findall(pattern, full_text)
        if not matches:
            return
        for match in matches:
            norm_key = normalize_key(match)
            raw_value = normalized_data.get(norm_key, "fyll i manuellt")
            if norm_key in suppress_keys:
                full_text = full_text.replace(f"{{{{{match}}}}}", "")
            else:
                translated = translate_value_if_possible(raw_value, key=norm_key)
                adjusted = adjust_translation_by_context(full_text, f"{{{{{match}}}}}", translated)
                full_text = full_text.replace(f"{{{{{match}}}}}", adjusted)
        # Töm alla runs
        for run in paragraph.runs:
            run.text = ""
        # Skriv tillbaka som en run
        if paragraph.runs:
            paragraph.runs[0].text = full_text
        else:
            paragraph.add_run(full_text)


    def process_all_paragraphs_in_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_runs(para)
                    process_all_paragraphs_in_tables(cell.tables)

    for para in doc.paragraphs:
        process_runs(para)
    process_all_paragraphs_in_tables(doc.tables)

    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            process_runs(para)
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_runs(para)

def fill_group_headings_dynamic(doc, group_headings):
    pattern = "{{section_heading}}"

    def process_table(table):
        for row in list(table.rows):
            for cell in row.cells:
                if pattern in cell.text:
                    template_row = row
                    table._tbl.remove(row._tr)
                    for heading in group_headings:
                        new_row = table.add_row()
                        for i, cell in enumerate(row.cells):
                            new_cell = new_row.cells[i]
                            text = cell.text.replace(pattern, heading)
                            new_cell.text = text
                    return True  # Avsluta efter att vi har ersatt
                for nested_table in cell.tables:
                    if process_table(nested_table):
                        return True
        return False

    for table in doc.tables:
        if process_table(table):
            break  # Klar efter första ersättningen

def fill_dynamic_text_rows(doc: Document, elevators: list, key, placeholder: str,
                           singular_template: str, grouped_template: str,
                           passive_template: str = None):
    from collections import defaultdict

    groups = defaultdict(list)

    for e in elevators:
        label = e.get("general_information", "").strip().split()[0].upper()

        if isinstance(key, tuple):  # Flera nycklar (t.ex. handrail_type och handrail_material)
            values = tuple(e.get(k, "").strip() for k in key)
            if all(values):
                groups[values].append(label)
            else:
                groups[None].append(label)
        else:
            value = e.get(key, "").strip()
            if value:
                groups[value].append(label)
            else:
                groups[None].append(label)

    texts = []

    def adjust(text_template, *translated_values):
        # Anpassa varje värde efter kontexten i mallen
        adjusted_values = []
        for value in translated_values:
            adjusted = adjust_translation_by_context(text_template, "{}", value)
            adjusted_values.append(adjusted)
        return text_template.format(*adjusted_values)

    # Fall 1: Alla hissar har samma (och det finns ett värde)
    if len(groups) == 1 and None not in groups:
        only_key = list(groups.keys())[0]
        if isinstance(only_key, tuple):
            translated_parts = [translate_value_if_possible(v, key=k) for v, k in zip(only_key, key)]
            texts.append(adjust(singular_template, *translated_parts))
        else:
            translated = translate_value_if_possible(only_key, key=key)
            texts.append(adjust(singular_template, translated))

    # Fall 2: Flera grupper
    else:
        for val, hissar in groups.items():
            hiss_text = ", ".join(sorted(hissar))
            if val is None and passive_template:
                texts.append(passive_template.format(hiss_text))
            elif val is not None:
                if isinstance(val, tuple):
                    translated_parts = [translate_value_if_possible(v, key=k) for v, k in zip(val, key)]
                    texts.append(adjust(grouped_template, hiss_text, *translated_parts))
                else:
                    translated = translate_value_if_possible(val, key=key)
                    texts.append(adjust(grouped_template, hiss_text, translated))

    def replace_placeholder_in_table(table):
        for row in list(table.rows):
            for cell in row.cells:
                if placeholder in cell.text:
                    table._tbl.remove(row._tr)
                    for text in texts:
                        new_row = table.add_row()
                        for i, c in enumerate(row.cells):
                            new_cell = new_row.cells[i]
                            if i == 0:
                                para = new_cell.paragraphs[0]
                                for run in para.runs:
                                    run.text = ""
                                para.add_run(text)

                                # Kopiera stil
                                original_para = cell.paragraphs[0]
                                para.style = original_para.style
                                para.paragraph_format.left_indent = original_para.paragraph_format.left_indent
                                para.paragraph_format.first_line_indent = original_para.paragraph_format.first_line_indent
                                para.paragraph_format.right_indent = original_para.paragraph_format.right_indent
                                para.paragraph_format.space_before = Pt(3)
                            else:
                                new_cell.text = c.text
                    return True
                for nested in cell.tables:
                    if replace_placeholder_in_table(nested):
                        return True
        return False

    for table in doc.tables:
        if replace_placeholder_in_table(table):
            break

def fill_static_row_if_present(doc: Document, elevators: list, key: str, value_filter, placeholder: str, text_template: str, fallback_text: str = None):
    """
    Om någon hiss uppfyller villkoret (via value_filter), ersätt placeholder med angiven text.
    text_template måste ha två {}: en för hissbeteckningar, en för värde.
    fallback_text används om ingen hiss matchar.
    """
    matched_labels = []
    matched_value = None

    for e in elevators:
        value = e.get(key, "").strip()
        if value_filter(value):
            label = e.get("general_information", "").strip().split()[0].upper()
            matched_labels.append(label)
            matched_value = value

    if matched_labels:
        hissar = ", ".join(sorted(set(matched_labels)))
        text = text_template.format(hissar, matched_value)
    elif fallback_text:
        text = fallback_text
    else:
        return  # Inget att fylla och ingen fallback

    def replace_placeholder_in_table(table):
        for row in list(table.rows):
            for cell in row.cells:
                if placeholder in cell.text:
                    table._tbl.remove(row._tr)
                    new_row = table.add_row()
                    for i, c in enumerate(row.cells):
                        new_cell = new_row.cells[i]
                        if i == 0:
                            para = new_cell.paragraphs[0]
                            for run in para.runs:
                                run.text = ""
                            para.add_run(text)

                            original_para = cell.paragraphs[0]
                            para.style = original_para.style
                            para.paragraph_format.left_indent = original_para.paragraph_format.left_indent
                            para.paragraph_format.first_line_indent = original_para.paragraph_format.first_line_indent
                            para.paragraph_format.right_indent = original_para.paragraph_format.right_indent
                            para.paragraph_format.space_before = Pt(3)  # lägg till spacing om du vill
                        else:
                            new_cell.text = c.text
                    return True
                for nested in cell.tables:
                    if replace_placeholder_in_table(nested):
                        return True
        return False

    for table in doc.tables:
        if replace_placeholder_in_table(table):
            break


def remove_rows_for_placeholders(doc: Document, placeholder_keys: list, data: dict):
    norm_keys = {normalize_key(k) for k in placeholder_keys}

    def remove_rows_in_table(table):
        for row in list(table.rows):
            row_text = " ".join(cell.text for cell in row.cells)

            for norm_key in norm_keys:
                if f"{{{{{norm_key}}}}}" in row_text:
                    value = data.get(norm_key, "")
                    if not value or str(value).strip() == "":
                        print(f"  Tar bort rad – nyckel saknas eller tom: {norm_key}")
                        row._tr.getparent().remove(row._tr)
                        break

            # Bearbeta nested tables även om raden togs bort
            for cell in row.cells:
                for nested_table in cell.tables:
                    remove_rows_in_table(nested_table)

    for table in doc.tables:
        remove_rows_in_table(table)

def remove_different_first_page(doc):
    for section in doc.sections:
        title_pg = section._sectPr.find(qn("w:titlePg"))
        if title_pg is not None:
            section._sectPr.remove(title_pg)

def insert_section_break_next_page(doc):
    try:
        doc.add_section(WD_SECTION.NEW_PAGE)
        print(" Section Break (Next Page) infogad korrekt.")
    except Exception as e:
        print(f" Kunde inte infoga Section Break: {e}")

def clear_headers_and_footers(doc):
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            for para in header.paragraphs:
                p = para._element
                p.getparent().remove(p)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            for para in footer.paragraphs:
                p = para._element
                p.getparent().remove(p)
    return doc

def remove_empty_paragraphs_before_first_table(doc):
    for p in list(doc.paragraphs):
        if not p.text.strip():
            p._element.getparent().remove(p._element)
        else:
            break

def remove_paragraphs_with_drawing_no_text_raw(doc):
    body = doc._element.body
    for p in list(body.iter(qn("w:p"))):
        has_drawing = p.find(".//" + qn("w:drawing")) is not None
        has_text = p.find(".//" + qn("w:t")) is not None
        if has_drawing and not has_text:
            p.getparent().remove(p)

def copy_margins_from_template(source, target):
    template_sec = source.sections[0]
    for section in target.sections:
        section.top_margin = template_sec.top_margin
        section.bottom_margin = template_sec.bottom_margin
        section.left_margin = template_sec.left_margin
        section.right_margin = template_sec.right_margin
        section.header_distance = template_sec.header_distance
        section.footer_distance = template_sec.footer_distance

def placeholders_missing_in_all_elevators(placeholders: list, elevators: list):
    norm_keys = {normalize_key(p) for p in placeholders}
    for elevator in elevators:
        for key in norm_keys:
            if key in elevator and str(elevator[key]).strip():
                return False  # Någon hiss har ett värde
    return True  # Alla hissar saknar alla värden

def merge_elevator_data(elevators, global_data=None):
    merged = {}
    for elevator in elevators:
        for k, v in elevator.items():
            norm_key = normalize_key(k)
            if norm_key not in merged or not merged[norm_key].strip():
                merged[norm_key] = v
    if global_data:
        merged.update(global_data)
    return merged

def merge_groups_by_key(groups, key):
    merged = defaultdict(list)
    for g in groups:
        k = g.get(key, "")
        hissar = g.get("hissbeteckning", "")
        if k and hissar:
            merged[k].extend(h.strip() for h in hissar.split(","))
    merged_list = []
    for k, hissar in merged.items():
        base = groups[0].copy()
        base["ceiling_type"] = k
        base["hissbeteckning"] = ", ".join(sorted(set(hissar)))
        base["antal_hissar"] = str(len(set(hissar)))
        merged_list.append(base)
    return merged_list

def group_elevators_by_keys(elevators, keys):
    grouped = defaultdict(list)
    for elevator in elevators:
        group_key = tuple(elevator.get(k, "") for k in keys)
        grouped[group_key].append(elevator)

    result = []
    for group in grouped.values():
        base = group[0].copy()
        base["hissbeteckning"] = ", ".join(
            e.get("general_information", "").split()[0].upper()
            for e in group if e.get("general_information")
        )
        base["antal_hissar"] = str(len(group))
        for k in keys:
            base[k] = group[0].get(k, "")
        result.append(base)
    return result

def generate_final_doc(template_path, elevator_groups, all_elevators, hissida_path, avslut_path, translation_dict, group_defs, global_data=None):
    master = Document(template_path)

    # Skapa dynamiska grupprubriker från group_defs
    group_headings = []
    for i, group in enumerate(group_defs):
        names = [e.strip().split()[0].upper() for e in group.get("hissbeteckning", "").split(",") if e.strip()]
        if names:
            if len(names) == 1:
                heading = f"Grupp {i+1}: {names[0]}"
            else:
                heading = f"Grupp {i+1}: {names[0]}–{names[-1]}"
            group_headings.append(heading)

    if global_data is None:
        global_data = {}
    global_data["datum"] = datetime.today().strftime("%Y-%m-%d")

    # Första hissen fylls i i huvudmallen
    combined_first = {**global_data, **elevator_groups[0]}
    placeholders_to_check = ["prl", "ebd_emergency_battery_drive"]
    if placeholders_missing_in_all_elevators(placeholders_to_check, all_elevators):
        print(" Tar bort rader – ingen hiss har värden för:",placeholders_to_check)
        remove_rows_for_placeholders(master,placeholders_to_check,combined_first)
    else:
        print(" Behåller rader – minst en hiss har värden för:",placeholders_to_check)

    fill_group_headings_dynamic(master, group_headings)
    fill_placeholders_in_doc(master, combined_first)
    print(" Fyller in grupprubriker:", group_headings)
    copy_margins_from_template(master, master)

    composer = Composer(master)

    for elevator in elevator_groups[1:]:
        doc = Document(hissida_path)
        remove_different_first_page(doc)
        remove_empty_paragraphs_before_first_table(doc)
        remove_paragraphs_with_drawing_no_text_raw(doc)
        clear_headers_and_footers(doc)
        combined_data = {**global_data, **elevator}
        fill_placeholders_in_doc(doc, combined_data)
        copy_margins_from_template(master, doc)
        insert_section_break_next_page(doc)
        composer.append(doc)

    # Avslutningsdelen → använd alla individuella hissar
    avslut = Document(avslut_path)
    remove_different_first_page(avslut)
    remove_empty_paragraphs_before_first_table(avslut)
    remove_paragraphs_with_drawing_no_text_raw(avslut)
    clear_headers_and_footers(avslut)

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="ceiling_type",
        placeholder="{{ceiling_type_group}}",
        singular_template='Tak i hisskorgar ska vara av typ {}.',
        grouped_template='Hissar med beteckning {} har tak i hisskorgen av typ {}.'
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="flooring_material",
        placeholder="{{floor_type_group}}",
        singular_template='Golv i hissar ska vara av typ {}.',
        grouped_template='Hissar med beteckning {} skall ha golv av typ {}.',
        passive_template="Hissar med beteckning {} skall ha lokalt golv av typ [fyll i vilket golv]."
    )


    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="car_door_panel_decoration_aside",
        placeholder="{{car_door_panel_group}}",
        singular_template='Korgdörrar skall vara av material {}.',
        grouped_template='Hissar med beteckning {} skall ha korgdörrar av typ {}.'
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="car_front_wall_material",
        placeholder="{{car_front_wall_material_group}}",
        singular_template='Korgöppningar för hiss skall vara av material {}.',
        grouped_template='Korgöppningar för hissar med beteckning {} skall vara av material {}.'
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="maximum_starts_per_hour",
        placeholder="{{maximum_starts_per_hour_group}}",
        singular_template='Drivsystem skall vara dimensionerat för minst {} starter per timma.',
        grouped_template='Drivsystem för hissar med beteckning {} skall vara dimensionerat för minst {} starter per timme.'
    )


    fill_static_row_if_present(
        avslut,
        all_elevators,
        key="elevator_complementary_standard",
        value_filter=lambda v: "EN81-72 2020" in v,
        placeholder="{{elevator_complementary_standard_group}}",
        text_template="Hissar med beteckning {} skall vara brandbekämpningshissar enligt {}.",
        fallback_text="Inga hissar är brandbekämpningshissar."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key=("handrail_type", "handrail_material"),
        placeholder="{{handrail_group}}",
        singular_template='Handledare skall vara {} i {}, på distans från korgvägg. Handledare monteras med överkant 900 mm över golv. Alla kanter, infästningar etc skall vara väl rundade och avfasade.',
        grouped_template='Hissar med beteckning {} skall ha handledare av typ {} i {}, på distans från korgvägg. Handledare monteras med överkant 900 mm över golv. Alla kanter, infästningar etc skall vara väl rundade och avfasade.'
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key=("car_mirror_size", "car_mirror_position"),
        placeholder="{{car_mirror_group}}",
        singular_template="Spegel skall vara {}. Spegel skall monteras på korgs {}.",
        grouped_template="Hissar med beteckning {} skall ha {} spegel, monterad på korgs {}.",
        passive_template="Hissar med beteckning {} saknar angiven typ för spegel."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="car_door_model_a_side",
        placeholder="{{car_door_model_a_side_group}}",
        singular_template="Korgdörr inklusive dörrmaskineri skall vara utförda och konstruerade för minst {} cykler (öppning och stängning) per år.",
        grouped_template="Hissar med beteckning {} skall ha korgdörr inklusive dörrmaskineri utfört och konstruerat för minst {} cykler (öppning och stängning) per år."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="car_fan_type",
        placeholder="{{car_fan_type_group}}",
        singular_template="Hisskorgar skall förses med {} som skall styras med 5 minuters frånslagsfördröjning.",
        grouped_template="Hissar med beteckning {} skall ha {}, styrd med 5 minuters frånslagsfördröjning.",
        passive_template="Hissar med beteckning {} skall förses med passiv ventilation i erforderlig omfattning."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="flip_chair_type",
        placeholder="{{flip_chair_type_group}}",
        singular_template="Fällsits skall monteras på korgvägg. Korgvägg skall förstärkas för infästning av fällsits.",
        grouped_template="Hissar med beteckning {} skall ha fällsits monterad på korgvägg. Korgvägg skall förstärkas för infästning av fällsits."
      )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="buffer_rails_quantity",
        placeholder="{{buffer_rails_quantity_group}}",
        singular_template="{} rad/rader med avbärarlister skall monteras ovan sockel på vägg som ej har dörröppning.",
        grouped_template="Hissar med beteckning {} skall ha {} rad/rader med avbärarlister monterad ovan sockel på vägg som ej har dörröppning.",
        passive_template="Hissar med beteckning {} skall ej ha avbärarlister."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="car_door_panel_decoration_aside",
        placeholder="{{car_door_panel_decoration_aside_group}}",
        singular_template="Korgdörrar skall vara av {}.",
        grouped_template="Hissar med beteckning {} skall ha korgdörrar av typ {}.",
        passive_template="Hissar med beteckning {} saknar angivet material för korgdörrar."
    )

    fill_static_row_if_present(
        avslut,
        all_elevators,
        key="prl",
        value_filter=lambda v: bool(v.strip()),
        placeholder="{{prl_group}}",
        text_template="Följande hissar skall ha prioriterad körning: {}."
    )

    fill_static_row_if_present(
        avslut,
        all_elevators,
        key="ebd_emergency_battery_drive",
        value_filter=lambda v: bool(v.strip()),
        placeholder="{{ebd_emergency_battery_drive_group}}",
        text_template="Automatisk nödsänkning krävs för följande hissar: {}."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="door_type",
        placeholder="{{door_type_group}}",
        singular_template="Schaktdörr skall vara av typ {} och med dagöppningar motsvarande korgdörrar.",
        grouped_template="Hissar med beteckning {} skall ha schaktdörr av typ {}, med dagöppningar motsvarande korgdörrar.",
        passive_template="Hissar med beteckning {} saknar angiven schaktdörrstyp."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="landing_door_model",
        placeholder="{{landing_door_model_group}}",
        singular_template="Schaktdörrar skall vara utförda och konstruerade för minst {} cykler (öppning och stängning) per år.",
        grouped_template="Schaktdörrar i hissar med beteckning {} skall vara utförda och konstruerade för minst {} cykler (öppning och stängning) per år."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="finishing_a",
        placeholder="{{finishing_a_group}}",
        singular_template="Schaktdörrar för hiss skall vara av typ {}.",
        grouped_template="Hissar med beteckning {} skall ha schaktdörrar av typ {}.",
        passive_template="Hissar med beteckning {} saknar angiven typ för schaktdörr."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key=("landing_door_frame_front", "finishing_a"),
        placeholder="{{landing_door_frame_front_group}}",  # Vi ersätter bara den ena, men båda behövs
        singular_template="För hiss monteras schaktdörrar i {} utförande, samt skall vara av material {}.",
        grouped_template="Hissar med beteckning {} skall ha schaktdörrar i {} utförande och av material {}.",
        passive_template="Hissar med beteckning {} saknar angivet dörrkarm- eller materialutförande."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key=("sill_type_a", "landing_door_sill_material"),
        placeholder="{{sill_type_a_group}}",  # Vi använder bara en placeholder för radmatchning
        singular_template="Trösklar skall utföras som {} med tröskelprofil i {}.",
        grouped_template="Hissar med beteckning {} skall ha trösklar utförda som {} med tröskelprofil i {}.",
        passive_template="Hissar med beteckning {} saknar uppgift om tröskeltyp eller material."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="signalisation_series",
        placeholder="{{signalisation_series_group}}",
        singular_template="Manöver- och indikeringsdon i hissarna skall vara av typ {}.",
        grouped_template="Hissar med beteckning {} skall ha manöver- och indikeringsdon av typ {}.",
        passive_template="Hissar med beteckning {} saknar angiven typ av manöver- och indikeringsdon."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key=("lcs_lci_material", "cop_face_plate_material"),
        placeholder="{{cover_plate_materials_group}}",
        singular_template="Täcklock för anrop skall vara av material {}, täcklock för tryckknappspanel i korg skall vara av material {}.",
        grouped_template="Hissar med beteckning {} skall ha täcklock för anrop av material {}, och täcklock för tryckknappspanel i korg av material {}.",
        passive_template="Hissar med beteckning {} saknar angivet material för täcklock vid anrop eller i korg."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="control_system",
        placeholder="{{control_system_group}}",
        singular_template="Hissar skall ha styrning för {}.",
        grouped_template="Hissar med beteckning {} skall ha styrning för {}.",
        passive_template="Hissar med beteckning {} saknar angiven styrningstyp."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="lcs_lci_placement",
        placeholder="{{lcs_lci_placement_group}}",
        singular_template="Anropsknapp, var med hiss kan kallas till stannplanet, skall placeras vid sidan av schaktdörr. Anropsknappar skall placeras i {}.",
        grouped_template="Hissar med beteckning {} skall ha anropsknapp, var med hiss kan kallas till stannplanet, placerad vid sidan av schaktdörr. Anropsknappar skall placeras i {}.",
        passive_template="Hissar med beteckning {} saknar angiven placering för anropsknappar."
    )

    fill_dynamic_text_rows(
        avslut,
        all_elevators,
        key="wall_b_finishing",
        placeholder="{{wall_b_finishing_group}}",
        singular_template="Korgväggar i hiss skall vara av typ {}.",
        grouped_template="Hissar med beteckning {} skall ha korgväggar av typ {}.",
        passive_template="Hissar med beteckning {} saknar angiven typ för korgväggar."
    )



    # Sammanfoga och fyll i resterande placeholders
    combined_all = merge_elevator_data(all_elevators, global_data)
    remove_rows_for_placeholders(avslut, placeholders_to_check, combined_all)
    fill_placeholders_in_doc(avslut, combined_all, suppress_keys=["prl", "ebd_emergency_battery_drive"])
    copy_margins_from_template(master, avslut)
    composer.append(avslut)

    try:
        for i in range(1, len(composer.doc.sections)):
            section = composer.doc.sections[i]
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True
    except Exception as e:
        print(f" Kunde inte länka sektioner till föregående: {e}")

    base_dir = os.path.dirname(os.path.abspath(__file__))  # Mapp där RFQ_GIT.py ligger
    output_dir = os.path.join(base_dir, "output")
    os.makedirs(output_dir, exist_ok=True)  # Skapar mappen om den inte finns

    output_path = os.path.join(output_dir, "komplett_rfqdokument.docx")

    total_pages = 7 + (len(elevator_groups) - 1) + 18
    for para in composer.doc.paragraphs:
        if "{{numpag}}" in para.text:
            for run in para.runs:
                run.text = run.text.replace("{{numpag}}", str(total_pages))

    composer.save(output_path)

    print(f" Dokument klart: Huvudmall + {len(elevator_groups)} hissidor + avslutningsmall.")
    return output_path

# --- Main ---

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(" Du måste ange XML-sökväg som argument.")
        sys.exit(1)

    xml_path = sys.argv[1]

    # Hårdkodade mallvägar (relativa till script)
    base_path = os.path.join(os.path.dirname(__file__), "backend_data")
    docx_file = os.path.join(base_path, "startmall lägga till rader.docx")
    translation_file = os.path.join(base_path, "database_RFQ.xlsx")
    hissida_path = os.path.join(base_path, "hissida.docx")
    avslut_path = os.path.join(base_path, "avslutningsmall dynamisk ta bort delar.docx")

    # Läs in översättningstabell
    df = pd.read_excel(translation_file)
    translation_dict = {
        normalize_key(str(row["english"]).strip()): str(row["generic_swedish"]).strip()
        for _, row in df.iterrows()
        if pd.notna(row["english"]) and pd.notna(row["generic_swedish"])
    }

    # Kör extraktion och grupperingar
    elevators, global_data = extract_multiple_elevators(xml_path)
    group_defs = extract_elevator_groups_from_xml(xml_path)
    elevator_groups = group_elevators_by_spec(elevators)

    # Skicka in den sammansatta datan till dokumentgenerering
    final_path = generate_final_doc(
        docx_file,
        elevator_groups,
        elevators,
        hissida_path,
        avslut_path,
        translation_dict,
        group_defs,
        global_data
    )

    # Flytta utdata till samma katalog som XML-filen finns i
    output_path = os.path.join(os.path.dirname(xml_path), "komplett_rfqdokument.docx")

    print(" Dokument genererat:", output_path)
